import os
import sys
import json
import yaml
from pathlib import Path
from docx import Document

import pytest

sys.path.append(os.path.join(os.path.dirname(__file__), ".."))
from scripts import patient_recruitment_material_agent as agent


class DummyMessage:
    def __init__(self, content: str):
        self.content = content


class DummyChoice:
    def __init__(self, content: str):
        self.message = DummyMessage(content)


class DummyResponse:
    def __init__(self, content: str):
        self.choices = [DummyChoice(content)]


def create_docx(path: Path, text: str) -> None:
    doc = Document()
    doc.add_paragraph(text)
    doc.save(path)


def test_generate_materials(tmp_path, mocker):
    protocol = tmp_path / "protocol.docx"
    insights = tmp_path / "insights.json"
    out_dir = tmp_path / "out"

    create_docx(protocol, "protocol text")
    insights.write_text(json.dumps({"patients": 50}))

    mocker.patch("scripts.patient_recruitment_material_agent.completion", return_value=DummyResponse("ad copy"))
    mocker.patch("scripts.patient_recruitment_material_agent.get_llm_model_name", return_value="model")

    paths = agent.generate_materials(str(protocol), str(insights), str(out_dir))

    assert Path(paths["html"]).exists()
    assert "ad copy" in Path(paths["html"]).read_text()

    assert Path(paths["docx"]).exists()
    doc = Document(paths["docx"])
    content = "\n".join(p.text for p in doc.paragraphs)
    assert "ad copy" in content

    assert Path(paths["png"]).exists()
    assert Path(paths["png"]).read_bytes() == b"PNG_PLACEHOLDER"


def test_update_checklist(tmp_path):
    checklist_data = [
        {"agentId": "2.200", "name": "Patient Recruitment Material Generator", "status": 0, "dependencies": []}
    ]
    file_path = tmp_path / "checklist.yml"
    file_path.write_text(yaml.safe_dump(checklist_data))

    agent.update_checklist(str(file_path), 88)

    updated = yaml.safe_load(file_path.read_text())
    assert updated[0]["status"] == 88


def test_write_progress_log(tmp_path):
    log_dir = tmp_path / "logs"
    path = agent.write_progress_log(str(log_dir), 20, "summary")

    log_path = Path(path)
    assert log_path.exists()
    data = json.loads(log_path.read_text())
    assert data["agentId"] == "2.200"
    assert data["status"] == 20
    assert data["summary"] == "summary"
    assert "timestamp" in data
