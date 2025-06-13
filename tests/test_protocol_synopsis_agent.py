import os
import sys
import json
import yaml
from pathlib import Path
import pytest
import datetime # Ensure datetime is imported
from docx import Document

# Add project root to the Python path
sys.path.append(os.path.join(os.path.dirname(__file__), ".."))
from scripts import protocol_synopsis_agent as agent

# --- ROBUST MOCKING CLASSES ---
# These helper classes more accurately mock the structure of a litellm response.
class MockLLMMessage:
    def __init__(self, content):
        self.content = content

class MockLLMChoice:
    def __init__(self, content):
        self.message = MockLLMMessage(content)

class MockLLMResponse:
    def __init__(self, content=""):
        self.choices = [MockLLMChoice(content)]

# --- TESTS ---

def test_generate_synopsis_creates_docx(tmp_path, mocker):
    # 1. Setup
    input_data = {
        'therapeuticArea': 'Oncology',
        'productName': 'ABC123',
        'studyPhase': 'Phase 2',
        'primaryObjective': 'Evaluate safety'
    }
    # Use the more robust mock response class
    mock_resp = MockLLMResponse('Rationale\nStudy Design')
    mocker.patch('litellm.completion', return_value=mock_resp)

    # 2. Execution
    out_dir = tmp_path / 'output'
    path = agent.generate_synopsis(input_data, str(out_dir))
    
    # 3. Assertion
    assert Path(path).exists()
    doc = Document(path)
    assert 'Rationale' in doc.paragraphs[0].text
    assert 'Study Design' in doc.paragraphs[1].text


def test_update_checklist(tmp_path):
    # 1. Setup
    checklist_data = [
        {'agentId': '1.100', 'name': 'Protocol Synopsis Generation Agent', 'status': 0, 'dependencies': []}
    ]
    file_path = tmp_path / 'checklist.yml'
    file_path.write_text(yaml.safe_dump(checklist_data))

    # 2. Execution
    agent.update_checklist(str(file_path), 80)

    # 3. Assertion
    updated_data = yaml.safe_load(file_path.read_text())
    assert updated_data[0]['status'] == 80


def test_write_progress_log_creates_valid_json(tmp_path):
    # 1. Setup
    log_dir = tmp_path / 'logs'
    test_status = 50
    test_summary = "Synopsis draft created."

    # 2. Execution
    path_str = agent.write_progress_log(str(log_dir), test_status, test_summary)
    
    # 3. Assertions
    log_path = Path(path_str)
    assert log_path.exists()
    assert log_path.name.startswith(f"1.100-{test_status}-")
    
    # Assert the content of the JSON file
    log_content = json.loads(log_path.read_text())
    assert log_content['agentId'] == '1.100'
    assert log_content['status'] == test_status
    assert log_content['summary'] == test_summary
    assert 'timestamp' in log_content
    assert len(log_content['timestamp']) > 10
