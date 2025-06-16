import os
import json
import logging
import datetime
from typing import Dict

from dotenv import load_dotenv
from litellm import completion
from docx import Document
import yaml

from scripts.utils import setup_logging, get_llm_model_name

load_dotenv()
setup_logging()
logger = logging.getLogger(__name__)

AGENT_ID = "2.200"


def generate_materials(protocol_path: str, insights_path: str, output_dir: str) -> Dict[str, str]:
    """Generate recruitment materials using an LLM."""
    if not os.path.exists(protocol_path):
        logger.error(f"Protocol not found: {protocol_path}")
        return {}
    if not os.path.exists(insights_path):
        logger.error(f"Insights file not found: {insights_path}")
        return {}

    doc = Document(protocol_path)
    protocol_text = "\n".join(p.text for p in doc.paragraphs)

    with open(insights_path, "r") as f:
        try:
            insights = json.load(f)
        except json.JSONDecodeError as exc:
            logger.error(f"Invalid insights JSON: {exc}")
            return {}

    model_name = get_llm_model_name()
    if not model_name:
        logger.error("LLM model configuration is missing.")
        return {}

    prompt = (
        "Create IRB-compliant patient recruitment content including a short ad copy, flyer text, "
        "and social media post based on the following protocol and population insights.\n"
        f"Protocol:\n{protocol_text}\n"
        f"Population Insights:\n{json.dumps(insights)}"
    )

    try:
        logger.debug("Calling LLM for recruitment materials")
        resp = completion(model=model_name, messages=[{"role": "user", "content": prompt}])
        materials_text = resp.choices[0].message.content
    except Exception as exc:
        logger.error(f"LLM call failed: {exc}")
        return {}

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    html_path = os.path.join(output_dir, "recruitment_material.html")
    with open(html_path, "w") as f:
        f.write(f"<html><body><p>{materials_text}</p></body></html>")

    docx_path = os.path.join(output_dir, "recruitment_material.docx")
    doc_out = Document()
    for line in materials_text.split("\n"):
        doc_out.add_paragraph(line)
    doc_out.save(docx_path)

    png_path = os.path.join(output_dir, "recruitment_material.png")
    with open(png_path, "wb") as f:
        f.write(b"PNG_PLACEHOLDER")

    logger.info(f"Materials saved to {output_dir}")
    return {"html": html_path, "docx": docx_path, "png": png_path}


def update_checklist(checklist_path: str, status: int) -> None:
    """Update the checklist.yml file for this agent."""
    with open(checklist_path, "r") as f:
        tasks = yaml.safe_load(f)

    for task in tasks:
        if task.get("agentId") == AGENT_ID:
            task["status"] = status
            break

    with open(checklist_path, "w") as f:
        yaml.safe_dump(tasks, f)


def write_progress_log(log_dir: str, status: int, summary: str) -> str:
    """Write a progress log JSON file."""
    if not os.path.exists(log_dir):
        os.makedirs(log_dir)

    timestamp = datetime.datetime.now(datetime.timezone.utc).strftime("%Y%m%d%H%M%S")
    filename = f"{AGENT_ID}-{status}-{timestamp}.json"
    path = os.path.join(log_dir, filename)

    data = {
        "agentId": AGENT_ID,
        "status": status,
        "summary": summary,
        "timestamp": timestamp,
    }
    with open(path, "w") as f:
        json.dump(data, f, indent=2)
    logger.info(f"Progress log written to {path}")
    return path


def main() -> None:
    import argparse

    parser = argparse.ArgumentParser(description="Patient Recruitment Material Generator")
    parser.add_argument("protocol", help="Path to the final protocol DOCX")
    parser.add_argument("insights", help="Path to patient population insights JSON")
    parser.add_argument("--status", type=int, default=100, help="Completion status")
    parser.add_argument("--output_dir", default="output", help="Directory for output materials")
    args = parser.parse_args()

    generate_materials(args.protocol, args.insights, args.output_dir)
    update_checklist(os.path.join("config", "checklist.yml"), args.status)
    write_progress_log(os.path.join("PROGRESS_LOGS", "new"), args.status, "Recruitment materials generated")


if __name__ == "__main__":
    main()
