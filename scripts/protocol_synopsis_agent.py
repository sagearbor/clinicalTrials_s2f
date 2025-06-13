import os
import json
import logging
from datetime import datetime
from typing import Dict

from dotenv import load_dotenv
from litellm import completion
from docx import Document
import yaml

from scripts.utils import setup_logging, get_llm_model_name

load_dotenv()
setup_logging()
logger = logging.getLogger(__name__)

AGENT_ID = "1.100"


def generate_synopsis(input_data: Dict[str, str], output_dir: str) -> str:
    """Generate a protocol synopsis DOCX using an LLM."""
    model_name = get_llm_model_name()
    if not model_name:
        logger.error("LLM model configuration is missing.")
        return ""

    prompt = (
        "Create a draft protocol synopsis with the following sections: "
        "Rationale, Study Design, Endpoints, and Statistical Methods. "
        f"Therapeutic Area: {input_data.get('therapeuticArea')}\n"
        f"Product Name: {input_data.get('productName')}\n"
        f"Study Phase: {input_data.get('studyPhase')}\n"
        f"Primary Objective: {input_data.get('primaryObjective')}"
    )

    try:
        logger.debug("Calling LLM for synopsis generation")
        resp = completion(model=model_name, messages=[{"role": "user", "content": prompt}])
        synopsis_text = resp.choices[0].message.content
    except Exception as exc:
        logger.error(f"LLM call failed: {exc}")
        return ""

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    doc = Document()
    doc.add_heading("Protocol Synopsis", level=1)
    for line in synopsis_text.split("\n"):
        doc.add_paragraph(line)

    output_path = os.path.join(output_dir, "protocol_synopsis.docx")
    doc.save(output_path)
    logger.info(f"Synopsis saved to {output_path}")
    return output_path


def update_checklist(checklist_path: str, status: int) -> None:
    """Update the checklist.yml file for this agent."""
    with open(checklist_path, "r") as f:
        tasks = yaml.safe_load(f)

    for task in tasks:
        if task.get("agentId") == AGENT_ID:
            task["status"] = status
            break

    with open(checklist_path, "w") as f:
        yaml.safe_dump(tasks, f)


def write_progress_log(log_dir: str, status: int, summary: str) -> str:
    """Write a progress log JSON file."""
    if not os.path.exists(log_dir):
        os.makedirs(log_dir)

    timestamp = datetime.utcnow().strftime("%Y%m%d%H%M%S")
    filename = f"{AGENT_ID}-{status}-{timestamp}.json"
    path = os.path.join(log_dir, filename)

    data = {
        "agentId": AGENT_ID,
        "status": status,
        "summary": summary,
        "timestamp": timestamp,
    }
    with open(path, "w") as f:
        json.dump(data, f, indent=2)
    logger.info(f"Progress log written to {path}")
    return path


def main():
    import argparse

    parser = argparse.ArgumentParser(description="Protocol Synopsis Generation Agent")
    parser.add_argument("input_json", help="Path to JSON file with study details")
    parser.add_argument("--status", type=int, default=100, help="Completion status")
    parser.add_argument("--output_dir", default="output", help="Directory for output DOCX")
    args = parser.parse_args()

    with open(args.input_json, "r") as f:
        input_data = json.load(f)

    generate_synopsis(input_data, args.output_dir)
    update_checklist(os.path.join("config", "checklist.yml"), args.status)
    write_progress_log(os.path.join("PROGRESS_LOGS", "new"), args.status, "Protocol synopsis generated")


if __name__ == "__main__":
    main()

