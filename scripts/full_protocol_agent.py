import os
import json
import logging
from datetime import datetime
from typing import Optional

from dotenv import load_dotenv
from litellm import completion
from docx import Document
import yaml

from scripts.utils import setup_logging, get_llm_model_name

load_dotenv()
setup_logging()
logger = logging.getLogger(__name__)

AGENT_ID = "1.400"


def generate_full_protocol(synopsis_path: str, template_path: Optional[str], output_dir: str) -> str:
    """Generate a full clinical trial protocol DOCX using an LLM."""
    model_name = get_llm_model_name()
    if not model_name:
        logger.error("LLM model configuration is missing.")
        return ""

    if not os.path.exists(synopsis_path):
        logger.error(f"Synopsis file not found: {synopsis_path}")
        return ""

    doc_syn = Document(synopsis_path)
    synopsis_text = "\n".join(p.text for p in doc_syn.paragraphs)

    prompt = (
        "Expand the following protocol synopsis into a full ICH E6-compliant clinical trial protocol.\n"
        f"Synopsis:\n{synopsis_text}\n"
    )

    if template_path and os.path.exists(template_path):
        doc_tpl = Document(template_path)
        template_sections = "\n".join(p.text for p in doc_tpl.paragraphs)
        prompt += f"\nUse the following template sections:\n{template_sections}\n"

    try:
        logger.debug("Calling LLM for full protocol generation")
        resp = completion(model=model_name, messages=[{"role": "user", "content": prompt}])
        protocol_text = resp.choices[0].message.content
    except Exception as exc:
        logger.error(f"LLM call failed: {exc}")
        return ""

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    doc = Document()
    doc.add_heading("Full Clinical Trial Protocol", level=1)
    for line in protocol_text.split("\n"):
        doc.add_paragraph(line)

    output_path = os.path.join(output_dir, "full_protocol.docx")
    doc.save(output_path)
    logger.info(f"Full protocol saved to {output_path}")
    return output_path


def update_checklist(checklist_path: str, status: int) -> None:
    """Update the checklist.yml file for this agent."""
    with open(checklist_path, "r") as f:
        tasks = yaml.safe_load(f)

    for task in tasks:
        if task.get("agentId") == AGENT_ID:
            task["status"] = status
            break

    with open(checklist_path, "w") as f:
        yaml.safe_dump(tasks, f)


def write_progress_log(log_dir: str, status: int, summary: str) -> str:
    """Write a progress log JSON file."""
    if not os.path.exists(log_dir):
        os.makedirs(log_dir)

    timestamp = datetime.now(datetime.UTC)().strftime("%Y%m%d%H%M%S")
    filename = f"{AGENT_ID}-{status}-{timestamp}.json"
    path = os.path.join(log_dir, filename)

    data = {
        "agentId": AGENT_ID,
        "status": status,
        "summary": summary,
        "timestamp": timestamp,
    }
    with open(path, "w") as f:
        json.dump(data, f, indent=2)
    logger.info(f"Progress log written to {path}")
    return path


def main():
    import argparse

    parser = argparse.ArgumentParser(description="Full Protocol Generation Agent")
    parser.add_argument("synopsis_docx", help="Path to the approved protocol synopsis DOCX")
    parser.add_argument("--template", help="Path to a protocol template DOCX", default="")
    parser.add_argument("--status", type=int, default=100, help="Completion status")
    parser.add_argument("--output_dir", default="output", help="Directory for output DOCX")
    args = parser.parse_args()

    generate_full_protocol(args.synopsis_docx, args.template or None, args.output_dir)
    update_checklist(os.path.join("config", "checklist.yml"), args.status)
    write_progress_log(os.path.join("PROGRESS_LOGS", "new"), args.status, "Full protocol generated")


if __name__ == "__main__":
    main()
